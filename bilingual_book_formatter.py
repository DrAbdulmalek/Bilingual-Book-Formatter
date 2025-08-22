#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Bilingual Book Formatter v2.3
أداة متطورة لمعالجة الكتب والمستندات ثنائية اللغة
"""

import json
import os
import sys
import tempfile
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import multiprocessing as mp
from concurrent.futures import ThreadPoolExecutor
import argparse

# Document processing imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
except ImportError:
    print("تحذير: مكتبة python-docx غير مثبتة")

try:
    import pdfplumber
    import pypdfium2 as pdfium
except ImportError:
    print("تحذير: مكتبات PDF غير مثبتة")

try:
    import ebooklib
    from ebooklib import epub
except ImportError:
    print("تحذير: مكتبة ebooklib غير مثبتة")

try:
    from PIL import Image, ImageOps
except ImportError:
    print("تحذير: مكتبة Pillow غير مثبتة")

# Translation imports
try:
    import deepl
except ImportError:
    print("تحذير: مكتبة DeepL غير مثبتة")

# Google Drive imports
try:
    from googleapiclient.discovery import build
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.http import MediaFileUpload
except ImportError:
    print("تحذير: مكتبات Google API غير مثبتة")

# GUI imports
try:
    from PyQt6.QtWidgets import (QApplication, QMainWindow, QVBoxLayout, QHBoxLayout,
                                 QWidget, QPushButton, QLabel, QFileDialog, QTextEdit,
                                 QProgressBar, QComboBox, QCheckBox, QSpinBox, QTabWidget,
                                 QScrollArea, QSplitter, QGroupBox, QGridLayout, QMessageBox)
    from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer
    from PyQt6.QtGui import QFont, QPixmap, QIcon
    GUI_AVAILABLE = True
except ImportError:
    GUI_AVAILABLE = False
    print("تحذير: مكتبة PyQt6 غير مثبتة - الواجهة الرسومية غير متاحة")


class BilingualBookFormatter:
    """فئة رئيسية لمعالجة الكتب ثنائية اللغة"""
    
    def __init__(self, config_path: str = "config.json"):
        """تهيئة المعالج مع ملف الإعدادات"""
        self.config = self.load_config(config_path)
        self.setup_logging()
        self.deepl_translator = None
        self.drive_service = None
        
        # تهيئة مترجم DeepL إذا كان متاحاً
        if self.config.get("translation", {}).get("enable_deepl", False):
            self.init_deepl()
        
        # تهيئة خدمة Google Drive إذا كانت متاحة
        self.init_google_drive()
    
    def load_config(self, config_path: str) -> Dict[str, Any]:
        """تحميل ملف الإعدادات"""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            logging.warning(f"ملف الإعدادات {config_path} غير موجود، استخدام الإعدادات الافتراضية")
            return self.get_default_config()
        except json.JSONDecodeError as e:
            logging.error(f"خطأ في تحليل ملف الإعدادات: {e}")
            return self.get_default_config()
    
    def get_default_config(self) -> Dict[str, Any]:
        """الإعدادات الافتراضية"""
        return {
            "page_margins": {"top": 2.5, "bottom": 2.5, "left": 2.0, "right": 2.0},
            "fonts": {
                "english": "Times New Roman",
                "arabic": "Traditional Arabic",
                "french": "Times New Roman",
                "chinese": "Noto Serif CJK SC",
                "persian": "Amiri",
                "hebrew": "DejaVu Sans"
            },
            "rtl_languages": ["arabic", "persian", "hebrew"],
            "styles": {
                "normal": {"size": 12, "bold": False, "color": "000000"}
            },
            "export_pdf": True,
            "export_epub": True,
            "translation": {
                "enable_deepl": False,
                "deepl_api_key": ""
            },
            "image_processing": {
                "enable": True,
                "image_position": "center",
                "max_width": 600,
                "format": "webp"
            }
        }
    
    def setup_logging(self):
        """إعداد نظام السجلات"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('bilingual_formatter.log', encoding='utf-8'),
                logging.StreamHandler(sys.stdout)
            ]
        )
    
    def init_deepl(self):
        """تهيئة مترجم DeepL"""
        try:
            api_key = self.config.get("translation", {}).get("deepl_api_key", "")
            if api_key:
                self.deepl_translator = deepl.Translator(api_key)
                logging.info("تم تهيئة مترجم DeepL بنجاح")
        except Exception as e:
            logging.error(f"فشل في تهيئة مترجم DeepL: {e}")
    
    def init_google_drive(self):
        """تهيئة خدمة Google Drive"""
        try:
            # هذا مجرد مثال - يحتاج إلى ملف credentials.json
            SCOPES = ['https://www.googleapis.com/auth/drive.file']
            creds = None
            
            if os.path.exists('token.json'):
                creds = Credentials.from_authorized_user_file('token.json', SCOPES)
            
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    if os.path.exists('credentials.json'):
                        flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
                        creds = flow.run_local_server(port=0)
                
                with open('token.json', 'w') as token:
                    token.write(creds.to_json())
            
            self.drive_service = build('drive', 'v3', credentials=creds)
            logging.info("تم تهيئة خدمة Google Drive بنجاح")
        except Exception as e:
            logging.warning(f"لم يتم تهيئة Google Drive: {e}")
    
    def extract_text_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """استخراج النص من ملف DOCX"""
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append({
                        'type': 'paragraph',
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # استخراج الصور
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    content.append({
                        'type': 'image',
                        'data': rel.target_part.blob,
                        'filename': rel.target_ref
                    })
            
            return content
        except Exception as e:
            logging.error(f"خطأ في استخراج النص من DOCX: {e}")
            return []
    
    def extract_text_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """استخراج النص من ملف PDF"""
        try:
            content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        content.append({
                            'type': 'paragraph',
                            'text': text,
                            'page': page_num + 1
                        })
            
            return content
        except Exception as e:
            logging.error(f"خطأ في استخراج النص من PDF: {e}")
            return []
    
    def extract_text_from_epub(self, file_path: str) -> List[Dict[str, Any]]:
        """استخراج النص من ملف EPUB"""
        try:
            book = epub.read_epub(file_path)
            content = []
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # تحليل HTML وإستخراج النص
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text = soup.get_text()
                    if text.strip():
                        content.append({
                            'type': 'paragraph',
                            'text': text,
                            'chapter': item.get_name()
                        })
            
            return content
        except Exception as e:
            logging.error(f"خطأ في استخراج النص من EPUB: {e}")
            return []
    
    def process_images(self, images: List[bytes]) -> List[str]:
        """معالجة الصور وتحويلها"""
        processed_images = []
        
        for i, image_data in enumerate(images):
            try:
                # تحويل إلى PIL Image
                image = Image.open(io.BytesIO(image_data))
                
                # تحسين الصورة
                image = ImageOps.exif_transpose(image)
                
                # تغيير الحجم إذا لزم الأمر
                max_width = self.config.get("image_processing", {}).get("max_width", 600)
                if image.width > max_width:
                    ratio = max_width / image.width
                    new_height = int(image.height * ratio)
                    image = image.resize((max_width, new_height), Image.Resampling.LANCZOS)
                
                # حفظ الصورة
                output_path = f"temp_image_{i}.{self.config.get('image_processing', {}).get('format', 'webp')}"
                image.save(output_path, optimize=True, quality=85)
                processed_images.append(output_path)
                
            except Exception as e:
                logging.error(f"خطأ في معالجة الصورة {i}: {e}")
        
        return processed_images
    
    def align_content(self, content1: List[Dict], content2: List[Dict]) -> List[Tuple[Dict, Dict]]:
        """محاذاة المحتوى بين اللغتين"""
        # خوارزمية بسيطة للمحاذاة - يمكن تحسينها
        aligned = []
        min_len = min(len(content1), len(content2))
        
        for i in range(min_len):
            aligned.append((content1[i], content2[i]))
        
        # إضافة العناصر المتبقية
        if len(content1) > min_len:
            for i in range(min_len, len(content1)):
                aligned.append((content1[i], None))
        elif len(content2) > min_len:
            for i in range(min_len, len(content2)):
                aligned.append((None, content2[i]))
        
        return aligned
    
    def create_docx_output(self, aligned_content: List[Tuple], output_path: str):
        """إنشاء مخرجات DOCX"""
        try:
            doc = Document()
            
            # إعداد الهوامش
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(self.config["page_margins"]["top"])
                section.bottom_margin = Inches(self.config["page_margins"]["bottom"])
                section.left_margin = Inches(self.config["page_margins"]["left"])
                section.right_margin = Inches(self.config["page_margins"]["right"])
            
            # إنشاء جدول بعمودين
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            for content1, content2 in aligned_content:
                row = table.add_row()
                
                # العمود الأول (اللغة الأولى)
                if content1:
                    if content1['type'] == 'paragraph':
                        row.cells[0].text = content1['text']
                
                # العمود الثاني (اللغة الثانية)
                if content2:
                    if content2['type'] == 'paragraph':
                        row.cells[1].text = content2['text']
                        # تطبيق محاذاة RTL إذا لزم الأمر
                        for paragraph in row.cells[1].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.save(output_path)
            logging.info(f"تم حفظ ملف DOCX: {output_path}")
            
        except Exception as e:
            logging.error(f"خطأ في إنشاء ملف DOCX: {e}")
    
    def create_epub_output(self, aligned_content: List[Tuple], output_path: str):
        """إنشاء مخرجات EPUB"""
        try:
            book = epub.EpubBook()
            book.set_identifier('bilingual_book')
            book.set_title('Bilingual Book')
            book.set_language('en')
            book.add_author('Bilingual Book Formatter')
            
            # إنشاء فصل
            chapter = epub.EpubHtml(title='Chapter 1', file_name='chap_01.xhtml', lang='en')
            
            html_content = '''
            <html>
            <head>
                <title>Bilingual Book</title>
                <style>
                    .bilingual-container { display: flex; }
                    .lang1 { flex: 1; padding: 10px; }
                    .lang2 { flex: 1; padding: 10px; text-align: right; }
                </style>
            </head>
            <body>
            '''
            
            for content1, content2 in aligned_content:
                html_content += '<div class="bilingual-container">'
                
                if content1 and content1['type'] == 'paragraph':
                    html_content += f'<div class="lang1">{content1["text"]}</div>'
                else:
                    html_content += '<div class="lang1"></div>'
                
                if content2 and content2['type'] == 'paragraph':
                    html_content += f'<div class="lang2">{content2["text"]}</div>'
                else:
                    html_content += '<div class="lang2"></div>'
                
                html_content += '</div>'
            
            html_content += '</body></html>'
            
            chapter.content = html_content
            book.add_item(chapter)
            
            # إضافة فهرس
            book.toc = (epub.Link("chap_01.xhtml", "Chapter 1", "intro"),)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # كتابة الكتاب
            epub.write_epub(output_path, book)
            logging.info(f"تم حفظ ملف EPUB: {output_path}")
            
        except Exception as e:
            logging.error(f"خطأ في إنشاء ملف EPUB: {e}")
    
    def process_books(self, lang1_path: str, lang2_path: str, output_base: str):
        """معالجة الكتب الرئيسية"""
        try:
            logging.info(f"بدء معالجة {lang1_path} و {lang2_path}")
            
            # استخراج المحتوى
            content1 = self.extract_content(lang1_path)
            content2 = self.extract_content(lang2_path)
            
            if not content1 or not content2:
                raise ValueError("فشل في استخراج المحتوى من أحد الملفات")
            
            # محاذاة المحتوى
            aligned_content = self.align_content(content1, content2)
            
            # إنشاء المخرجات
            if self.config.get("export_pdf", True):
                self.create_docx_output(aligned_content, f"{output_base}.docx")
            
            if self.config.get("export_epub", True):
                self.create_epub_output(aligned_content, f"{output_base}.epub")
            
            logging.info("تمت المعالجة بنجاح")
            
        except Exception as e:
            logging.error(f"خطأ في معالجة الكتب: {e}")
            raise
    
    def extract_content(self, file_path: str) -> List[Dict[str, Any]]:
        """استخراج المحتوى حسب نوع الملف"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.epub':
            return self.extract_text_from_epub(file_path)
        else:
            raise ValueError(f"نوع ملف غير مدعوم: {file_ext}")
    
    def upload_to_drive(self, file_path: str) -> str:
        """رفع ملف إلى Google Drive"""
        if not self.drive_service:
            raise ValueError("خدمة Google Drive غير مهيأة")
        
        try:
            file_metadata = {'name': os.path.basename(file_path)}
            media = MediaFileUpload(file_path)
            
            file = self.drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()
            
            logging.info(f"تم رفع الملف إلى Google Drive: {file.get('id')}")
            return file.get('id')
            
        except Exception as e:
            logging.error(f"خطأ في رفع الملف إلى Google Drive: {e}")
            raise


class BilingualBookFormatterGUI(QMainWindow):
    """واجهة رسومية لمعالج الكتب ثنائية اللغة"""
    
    def __init__(self):
        super().__init__()
        self.formatter = BilingualBookFormatter()
        self.init_ui()
    
    def init_ui(self):
        """تهيئة الواجهة الرسومية"""
        self.setWindowTitle("Bilingual Book Formatter v2.3")
        self.setGeometry(100, 100, 1000, 700)
        
        # الأدوات الرئيسية
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # التخطيط الرئيسي
        main_layout = QVBoxLayout(central_widget)
        
        # شريط علوي للملفات
        file_layout = QHBoxLayout()
        
        # اختيار الملف الأول
        self.lang1_label = QLabel("الملف الأول:")
        self.lang1_path = QLabel("لم يتم اختيار ملف")
        self.lang1_button = QPushButton("اختيار ملف")
        self.lang1_button.clicked.connect(self.select_lang1_file)
        
        file_layout.addWidget(self.lang1_label)
        file_layout.addWidget(self.lang1_path)
        file_layout.addWidget(self.lang1_button)
        
        # اختيار الملف الثاني
        self.lang2_label = QLabel("الملف الثاني:")
        self.lang2_path = QLabel("لم يتم اختيار ملف")
        self.lang2_button = QPushButton("اختيار ملف")
        self.lang2_button.clicked.connect(self.select_lang2_file)
        
        file_layout.addWidget(self.lang2_label)
        file_layout.addWidget(self.lang2_path)
        file_layout.addWidget(self.lang2_button)
        
        main_layout.addLayout(file_layout)
        
        # إعدادات المعالجة
        settings_group = QGroupBox("إعدادات المعالجة")
        settings_layout = QGridLayout(settings_group)
        
        # تنسيق الإخراج
        settings_layout.addWidget(QLabel("تنسيق الإخراج:"), 0, 0)
        self.output_format = QComboBox()
        self.output_format.addItems(["DOCX", "EPUB", "كلاهما"])
        settings_layout.addWidget(self.output_format, 0, 1)
        
        # معالجة الصور
        self.process_images_cb = QCheckBox("معالجة الصور")
        self.process_images_cb.setChecked(True)
        settings_layout.addWidget(self.process_images_cb, 1, 0)
        
        # رفع إلى Google Drive
        self.upload_drive_cb = QCheckBox("رفع إلى Google Drive")
        settings_layout.addWidget(self.upload_drive_cb, 1, 1)
        
        main_layout.addWidget(settings_group)
        
        # منطقة المعاينة
        preview_group = QGroupBox("معاينة المحتوى")
        preview_layout = QVBoxLayout(preview_group)
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        preview_layout.addWidget(self.preview_text)
        
        main_layout.addWidget(preview_group)
        
        # شريط التقدم
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
        # أزرار التحكم
        button_layout = QHBoxLayout()
        
        self.preview_button = QPushButton("معاينة")
        self.preview_button.clicked.connect(self.preview_content)
        button_layout.addWidget(self.preview_button)
        
        self.process_button = QPushButton("معالجة")
        self.process_button.clicked.connect(self.process_files)
        button_layout.addWidget(self.process_button)
        
        self.clear_button = QPushButton("مسح")
        self.clear_button.clicked.connect(self.clear_all)
        button_layout.addWidget(self.clear_button)
        
        main_layout.addLayout(button_layout)
        
        # متغيرات الملفات
        self.lang1_file_path = None
        self.lang2_file_path = None
    
    def select_lang1_file(self):
        """اختيار الملف الأول"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "اختيار الملف الأول", "",
            "المستندات المدعومة (*.docx *.pdf *.epub);;جميع الملفات (*)"
        )
        if file_path:
            self.lang1_file_path = file_path
            self.lang1_path.setText(os.path.basename(file_path))
    
    def select_lang2_file(self):
        """اختيار الملف الثاني"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "اختيار الملف الثاني", "",
            "المستندات المدعومة (*.docx *.pdf *.epub);;جميع الملفات (*)"
        )
        if file_path:
            self.lang2_file_path = file_path
            self.lang2_path.setText(os.path.basename(file_path))
    
    def preview_content(self):
        """معاينة المحتوى"""
        if not self.lang1_file_path or not self.lang2_file_path:
            QMessageBox.warning(self, "تحذير", "يرجى اختيار كلا الملفين أولاً")
            return
        
        try:
            # استخراج عينة من المحتوى
            content1 = self.formatter.extract_content(self.lang1_file_path)[:5]  # أول 5 فقرات
            content2 = self.formatter.extract_content(self.lang2_file_path)[:5]
            
            preview_text = "معاينة المحتوى:\n\n"
            
            for i, (c1, c2) in enumerate(zip(content1, content2)):
                preview_text += f"--- الفقرة {i+1} ---\n"
                if c1 and c1['type'] == 'paragraph':
                    preview_text += f"الملف الأول: {c1['text'][:100]}...\n"
                if c2 and c2['type'] == 'paragraph':
                    preview_text += f"الملف الثاني: {c2['text'][:100]}...\n"
                preview_text += "\n"
            
            self.preview_text.setText(preview_text)
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"خطأ في معاينة المحتوى: {str(e)}")
    
    def process_files(self):
        """معالجة الملفات"""
        if not self.lang1_file_path or not self.lang2_file_path:
            QMessageBox.warning(self, "تحذير", "يرجى اختيار كلا الملفين أولاً")
            return
        
        # اختيار مجلد الحفظ
        output_dir = QFileDialog.getExistingDirectory(self, "اختيار مجلد الحفظ")
        if not output_dir:
            return
        
        try:
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            
            # تحديد اسم الملف الناتج
            output_base = os.path.join(output_dir, "bilingual_output")
            
            # معالجة الملفات
            self.formatter.process_books(self.lang1_file_path, self.lang2_file_path, output_base)
            
            self.progress_bar.setValue(100)
            
            # رفع إلى Google Drive إذا طُلب
            if self.upload_drive_cb.isChecked():
                try:
                    if os.path.exists(f"{output_base}.docx"):
                        self.formatter.upload_to_drive(f"{output_base}.docx")
                    if os.path.exists(f"{output_base}.epub"):
                        self.formatter.upload_to_drive(f"{output_base}.epub")
                except Exception as e:
                    QMessageBox.warning(self, "تحذير", f"فشل في الرفع إلى Google Drive: {str(e)}")
            
            QMessageBox.information(self, "نجح", f"تمت المعالجة بنجاح!\nالملفات محفوظة في: {output_dir}")
            
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"خطأ في المعالجة: {str(e)}")
        finally:
            self.progress_bar.setVisible(False)
    
    def clear_all(self):
        """مسح جميع البيانات"""
        self.lang1_file_path = None
        self.lang2_file_path = None
        self.lang1_path.setText("لم يتم اختيار ملف")
        self.lang2_path.setText("لم يتم اختيار ملف")
        self.preview_text.clear()


def main():
    """الدالة الرئيسية"""
    parser = argparse.ArgumentParser(description="Bilingual Book Formatter v2.3")
    parser.add_argument("--gui", action="store_true", help="تشغيل الواجهة الرسومية")
    parser.add_argument("--lang1", help="مسار الملف الأول")
    parser.add_argument("--lang2", help="مسار الملف الثاني")
    parser.add_argument("--output", help="مسار الملف الناتج (بدون امتداد)")
    parser.add_argument("--format", choices=["docx", "epub", "both"], default="both", help="تنسيق الإخراج")
    
    args = parser.parse_args()
    
    if args.gui:
        if not GUI_AVAILABLE:
            print("خطأ: مكتبة PyQt6 غير مثبتة. لا يمكن تشغيل الواجهة الرسومية.")
            sys.exit(1)
        
        app = QApplication(sys.argv)
        window = BilingualBookFormatterGUI()
        window.show()
        sys.exit(app.exec())
    
    elif args.lang1 and args.lang2 and args.output:
        formatter = BilingualBookFormatter()
        try:
            formatter.process_books(args.lang1, args.lang2, args.output)
            print(f"تمت المعالجة بنجاح! الملفات محفوظة في: {args.output}")
        except Exception as e:
            print(f"خطأ في المعالجة: {e}")
            sys.exit(1)
    
    else:
        parser.print_help()


if __name__ == "__main__":
    main()

